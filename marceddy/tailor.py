"""Resume tailoring — genuinely per-job, emitted as .docx (+ .txt) for ATS.

Tailoring never invents content. Per job it leads with the skills the posting
actually mentions, names the target role, and orders experience by relevance to
the role type — but every line comes from the user's true profile.
"""
from .models import slugify

ROLE_HINTS = [
    ("helpdesk", ["help desk", "helpdesk", "service desk", "tier 1", "tier 2",
                  "support specialist", "end user", "end-user"]),
    ("desktop",  ["desktop", "deskside", "field service", "imaging", "hardware", "technician"]),
    ("sysadmin", ["administrator", "sysadmin", "active directory", "windows server",
                  "infrastructure", "systems"]),
    ("network",  ["network", "dns", "dhcp", "tcp/ip", "switch", "router", "vpn"]),
    ("cloud",    ["cloud", "azure", "aws", "gcp", "devops", "docker", "ci/cd"]),
    ("developer", ["developer", "engineer", "python", "api", "software", "backend"]),
]


def matched_skills(job, profile):
    text = job.text_blob()
    return [s for s in profile["skills"] if s.lower() in text]


def role_type(job):
    blob = job.text_blob()
    best, score = "helpdesk", 0
    for name, kws in ROLE_HINTS:
        c = sum(1 for k in kws if k in blob)
        if c > score:
            best, score = name, c
    return best


_EXP_KEYS = {
    "helpdesk": ["support", "service desk", "ticket", "servicenow", "tier"],
    "desktop": ["desktop", "imaging", "hardware", "provision", "technician"],
    "sysadmin": ["active directory", "server", "infrastructure", "admin", "entra"],
    "network": ["network", "dns", "dhcp", "subnet", "vpn"],
    "cloud": ["cloud", "docker", "gcp", "azure", "ci/cd", "deploy"],
    "developer": ["python", "fastapi", "api", "saas", "build"],
}


def _experience_order(job, profile):
    rt = role_type(job)
    keys = _EXP_KEYS.get(rt, [])

    def rel(e):
        t = (e.get("title", "") + " " + e.get("org", "") + " " +
             " ".join(e.get("bullets", []))).lower()
        return sum(1 for k in keys if k in t)
    return sorted(profile["experience"], key=rel, reverse=True)


def render_resume(job, profile, policy):
    job_skills = matched_skills(job, profile)
    ordered = job_skills + [s for s in profile["skills"] if s not in job_skills]
    b = [("title", profile["name"])]
    contact = " | ".join(x for x in [profile.get("location", ""), profile.get("email", ""),
                                     profile.get("phone", "")] if x)
    b.append(("contact", contact))
    b.append(("h", "SUMMARY"))
    b.append(("p", "Targeting the %s role at %s. %s" % (job.title, job.company, profile["summary"])))
    if job_skills:
        b.append(("h", "KEY QUALIFICATIONS FOR THIS ROLE"))
        b.append(("p", ", ".join(job_skills[:12])))
    b.append(("h", "CERTIFICATIONS"))
    b += [("bullet", c) for c in profile["certifications"]]
    b.append(("h", "CORE SKILLS"))
    b.append(("p", ", ".join(ordered)))
    b.append(("h", "EXPERIENCE"))
    for e in _experience_order(job, profile):
        b.append(("role", "%s, %s (%s)" % (e["title"], e["org"], e["dates"])))
        b += [("bullet", x) for x in e["bullets"]]
    b.append(("h", "EDUCATION"))
    b += [("bullet", ed) for ed in profile["education"]]
    return b


def blocks_to_text(blocks):
    L = []
    for k, t in blocks:
        if k == "title":
            L.append(t)
        elif k == "contact":
            L += [t, ""]
        elif k == "h":
            L += ["", t]
        elif k == "role":
            L += ["", t]
        elif k == "bullet":
            L.append("  - " + t)
        else:
            L.append(t)
    return "\n".join(L).rstrip() + "\n"


def blocks_to_docx(blocks, path):
    from docx import Document
    doc = Document()
    for k, t in blocks:
        if k == "title":
            doc.add_heading(t, level=0)
        elif k == "h":
            doc.add_heading(t, level=1)
        elif k == "role":
            p = doc.add_paragraph()
            p.add_run(t).bold = True
        elif k == "bullet":
            doc.add_paragraph(t, style="List Bullet")
        else:
            doc.add_paragraph(t)
    doc.save(path)


def resume_basename(job):
    # <company>_<role>_<jobid>
    return "%s_%s_%s" % (job.company_slug, job.role_slug, job.job_id)


def write_resume(job, profile, policy, config):
    """Write the tailored resume as .docx (primary) + .txt (for text paste).
    Returns the .docx path (the proper application file)."""
    config.ensure_dirs()
    blocks = render_resume(job, profile, policy)
    base = config.resumes_dir / resume_basename(job)
    txt = str(base) + ".txt"
    docx = str(base) + ".docx"
    with open(txt, "w") as f:
        f.write(blocks_to_text(blocks))
    try:
        blocks_to_docx(blocks, docx)
        return docx
    except Exception:
        return txt
